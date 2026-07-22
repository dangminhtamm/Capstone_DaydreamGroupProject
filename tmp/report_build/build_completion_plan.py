from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "output/report/DayDreamer_Project_Completion_Plan_improved.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=45, start=75, bottom=45, end=75):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths, header_fill=GRAY_FILL):
    set_table_width(table, widths)
    set_table_borders(table)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7.8)
            if i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = DARK_BLUE


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table, widths, header_fill)
    return table


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.4)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, color, before, after in [
        ("Heading 1", 13, BLUE, 7, 3),
        ("Heading 2", 11, BLUE, 5, 2),
        ("Heading 3", 10, DARK_BLUE, 4, 1),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(9.2)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.14)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("DayDreamer - Project Completion Plan")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    add_footer(section)
    set_styles(doc)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RMIT Classification: Trusted")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("Project Completion Plan")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("The Second Brain: Grounded Personal Memory and Reflection Platform")
    r.font.size = Pt(13)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.add_run("Course: OENG1185 / COSC2503 Capstone Project Part B / Programming Project 2\n")
    p.add_run("Team: DayDreamer\n")
    p.add_run("Academic supervisor: Dr. Jeff Nijsse\n")
    p.add_run("Industry supervisor: Mr. Vo Hoang Phuc\n")
    p.add_run("Company: Tuturuuu")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.add_run("Team members\n").bold = True
    p.add_run("Dang Minh Tam (s3980087)\n")
    p.add_run("Nguyen Thanh Nhan (s4073629)\n")
    p.add_run("Nguyen Tan Thang (s3986344)\n")
    p.add_run("Tran Nguyen Quan (s3979298)\n")
    p.add_run("Duong Minh Duc Anh (s3618640)")

    doc.add_page_break()

    # Page/body content
    doc.add_heading("1. Updated Project Description", level=1)
    add_para(
        doc,
        "The Second Brain is an authenticated personal intelligence platform that converts diary entries, "
        "calendar context, attachments, and summaries into searchable memory. The current MVP is not only a "
        "journal: it stores user-authored context, indexes it asynchronously into memory chunks, retrieves "
        "relevant evidence using hybrid vector and lexical search, and returns source-grounded answers with "
        "citations, confidence, timing, and token analytics."
    )
    add_para(
        doc,
        "The scope has been refined since the original proposal. Instead of expanding equally into every Google "
        "service and every file type, Capstone B now prioritises a measurable core path: Diary -> Indexing Worker "
        "-> Memory Retrieval -> Cited Answer -> Summary/Reflection. This adjustment responds to supervisor and "
        "mentor feedback that the system must prove trust, retrieval accuracy, and latency before adding more "
        "integrations."
    )

    add_table(
        doc,
        ["Scope item", "Current decision", "Justification"],
        [
            ["Core diary memory", "In scope and implemented", "Main value path; mood/tags were added so emotional reflection can be evaluated."],
            ["Google Calendar", "In scope and partly live", "OAuth, token storage, event sync, diary linking, and memory indexing are implemented; Google test-user verification remains an operational dependency."],
            ["Attachment ingestion", "In scope for MVP testing", "Upload, Supabase Storage, text extraction path, worker indexing, and attachment citations exist; PDF/image/DOCX extraction still needs more acceptance testing."],
            ["Gmail and Google Contacts", "Deferred/future work", "Schema and planning exist, but the team decided not to expand beyond Calendar until the cited memory loop is stable."],
            ["AI optimisation", "High priority", "Focus is on retrieval filters, reranking, grounded fallback, cache control, and measurable evaluation metrics."],
        ],
        [1.35, 1.45, 3.7],
        LIGHT_BLUE,
    )

    doc.add_heading("2. Completed Work and Current Status", level=1)
    add_para(
        doc,
        "The project has moved from a concept into a working modular monorepo with a Next.js frontend, NestJS API, "
        "PostgreSQL/Supabase data layer, pgvector memory storage, and a Node worker for background jobs."
    )

    add_table(
        doc,
        ["Area", "Completed evidence from code", "Remaining gap"],
        [
            ["Diary and Timeline", "Authenticated diary CRUD, mood selector, tag input, timeline display, edit/delete, signed attachment URLs.", "Final UI polish and mobile walkthrough testing."],
            ["AI Memory", "Gemini embeddings, semantic chunking, hybrid retrieval, date/source filters, reranking, citation validation, debug trace, fast/deep/auto modes, fallback for quota or invalid JSON.", "Run benchmark dataset and record Recall@5, citation precision, no-answer accuracy, and p95 latency."],
            ["Calendar", "Server-side OAuth callback, token refresh update, encrypted token support, event upsert, sync endpoint, diary-event linking, calendar event indexing jobs.", "Google consent screen test users must be maintained; more live-account regression tests required."],
            ["Attachments", "Upload endpoint supports text/PDF/image/DOCX MIME types, Supabase Storage, Attachment rows, text extraction for plain text, worker extraction/indexing for non-text files, citation-capable chunks.", "Complete controlled PDF/image/DOCX acceptance tests and document extraction accuracy."],
            ["Summary/Reflection", "Daily, weekly, monthly, yearly summary pipeline; summaries are queued for memory indexing and shown in Summary UI.", "Tune prompts and evaluate summary usefulness with user tasks."],
            ["Observability/QA", "Search UI exposes confidence, answer mode, pipeline timing, token usage, retrieved chunks, cache/fallback status; seed scripts and worker drain support repeatable demo data.", "API Jest suite has one module-resolution failure; worker build config and web build must be rechecked outside sandbox constraints."],
        ],
        [1.1, 3.25, 2.15],
        GRAY_FILL,
    )

    add_para(
        doc,
        "Feedback response: the report and implementation now emphasise measurable evidence. The AI answer layer rejects unsupported names/dates, uses citation checks before returning model answers, and falls back to extractive evidence rather than inventing answers. The team also added seed data and explicit demo readiness checks so testing is repeatable."
    )
    add_para(
        doc,
        "Verification snapshot on 19 July 2026: AI tests pass 40/40, worker tests pass 2/2, web tests pass 4/4, and AI/API builds pass. API Jest still has one module-resolution suite failure; worker and web production builds need a final non-sandbox configuration check."
    )

    doc.add_heading("3. Planned Deliverables and Dates", level=1)
    add_para(
        doc,
        "The remaining plan is organised around five delivery weeks. The first three weeks match the team's current test plan: real attachment ingestion, Google Calendar integration, and AI optimisation."
    )
    add_table(
        doc,
        ["Week / dates", "Owner focus", "Deliverables", "Acceptance evidence"],
        [
            ["W1: 21-27 Jul", "Attachment ingestion", "Upload real PDF/image/DOCX/text files; verify Storage, extraction, outbox jobs, memory chunks, and attachment citations.", "At least 3 file types indexed; search answer cites attachment source; no dead-letter jobs."],
            ["W2: 28 Jul-3 Aug", "Google Calendar", "Retest OAuth with approved Google test users; sync real events; refresh token; link diary entries to events; index Calendar memories.", "Connected status true; events in DB; linked diary count > 0; Calendar query returns cited event."],
            ["W3: 4-10 Aug", "AI optimisation", "Run benchmark questions, refine filters/reranking/fallbacks, separate retrieval latency from generation latency, disable stale cache in tests.", "Recall@5, citation precision, no-answer accuracy and p95 retrieval latency recorded."],
            ["W4: 11-17 Aug", "Reflection and readiness", "Improve summary quality, demo readiness dashboard, seed data, error/loading states, and requeue controls.", "Create -> Search -> Summary flow passes with realistic data."],
            ["W5: 18-24 Aug", "Final demo and report", "User testing, bug fixing, deployment rehearsal, final written evidence and showcase script.", "Demo script completed twice; user testing results summarised; high risks closed or scoped."],
        ],
        [0.9, 1.15, 3.0, 1.55],
        LIGHT_BLUE,
    )

    doc.add_heading("4. Quantifiable Success Metrics", level=1)
    add_table(
        doc,
        ["Metric", "Target", "Feature measured", "Evaluation method"],
        [
            ["Retrieval Recall@5", ">= 80%", "AI Memory retrieval", "Use seeded benchmark questions with known expected source chunks."],
            ["Citation precision", ">= 85%", "Grounded answer/citation UI", "Manually verify whether each citation supports the claim shown."],
            ["No-answer accuracy", ">= 80%", "Hallucination control", "Ask unsupported questions and check that the system refuses instead of inventing."],
            ["p95 retrieval latency", "<= 500 ms excluding Gemini generation", "Memory search backend", "Record embed/retrieve/generate timing separately from Search analytics."],
            ["Indexing reliability", "0 dead-letter jobs after drain; >= 95% succeeded", "Worker and outbox pipeline", "Run seed/drain scripts and inspect readiness/outbox status."],
            ["Calendar sync success", ">= 90% of test events synced and searchable", "Google Calendar integration", "Compare Google Calendar test events with `calendar_events` and search citations."],
            ["User task completion", ">= 80% completion across 5 core tasks", "Frontend/product UX", "Test users complete create diary, attach file, connect Calendar, ask Search, read Summary."],
        ],
        [1.15, 1.0, 1.55, 2.9],
        GRAY_FILL,
    )

    doc.add_heading("5. Required Resources", level=1)
    add_table(
        doc,
        ["Resource", "Why it is required"],
        [
            ["Supabase Postgres + pgvector + Storage", "Stores authenticated users, diary entries, Calendar events, attachments, summaries, memory chunks, and vectors."],
            ["Gemini API quota", "Required for embeddings, chunk generation, answer generation, summary generation, and attachment text extraction for non-text files."],
            ["Google Cloud OAuth test users", "Required because the app is in Testing mode; unapproved accounts receive 403 access_denied."],
            ["Vercel/Render or equivalent hosting", "Needed for stable frontend/API demo and OAuth callback URLs."],
            ["Weekly mentor/team integration time", "Needed to close cross-module risks: API contracts, UI binding, worker indexing, and demo evidence."],
        ],
        [1.75, 4.75],
        LIGHT_BLUE,
    )

    doc.add_heading("6. Risk Management Plan", level=1)
    add_table(
        doc,
        ["Risk", "Priority", "Mitigation / workaround"],
        [
            ["LLM hallucination or weak citation grounding", "High", "Keep strict RAG prompt, citation validation, unsupported-name/date checks, low-confidence fallback, and visible source snippets."],
            ["Gemini quota/rate limits block live indexing or answering", "High", "Throttle worker jobs, cache completed answers, prepare seeded memory chunks, and maintain backup API key or lower-cost model for indexing."],
            ["Google OAuth app blocks unapproved testers", "High", "Add all team/test accounts to Google OAuth test users; verify redirect URI before demo; document reconnect process."],
            ["Async indexing jobs stuck or dead-lettered", "High", "Use outbox status/readiness checks, requeue controls, drain script, and demo checklist before each rehearsal."],
            ["Attachment extraction quality varies by file type", "Medium", "Accept plain text as baseline, test PDF/image/DOCX separately, and mark advanced OCR/DOCX extraction as controlled MVP scope."],
            ["Frontend/backend contract mismatch", "Medium", "Freeze DTOs for Diary, Search, Summary, Calendar; run API client tests; use mock data only as fallback demo mode."],
            ["Documentation overclaims implementation", "Medium", "Keep report aligned with code evidence and explicitly separate completed, partial, and future work."],
        ],
        [1.75, 0.75, 4.0],
        GRAY_FILL,
    )

    add_para(
        doc,
        "Priority rule: if a high-risk item remains unresolved in any week, the team will stabilise the core MVP path first and move optional scope such as Gmail, Google Contacts, advanced OCR, and data export into future work."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
